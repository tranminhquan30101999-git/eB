import asyncio
from pathlib import Path
from typing import Optional
import aiofiles

async def process_uploaded_file(file_path: Path, file_extension: str) -> str:
    """Process uploaded file and extract text content"""
    
    try:
        if file_extension == ".txt":
            return await process_text_file(file_path)
        elif file_extension == ".pdf":
            return await process_pdf_file(file_path)
        elif file_extension in [".docx", ".doc"]:
            return await process_word_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    except Exception as e:
        raise Exception(f"Error processing file {file_path}: {str(e)}")

async def process_text_file(file_path: Path) -> str:
    """Process .txt files"""
    async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
        content = await f.read()
    return content.strip()

async def process_pdf_file(file_path: Path) -> str:
    """Process .pdf files using PyPDF2"""
    try:
        import PyPDF2
        
        def extract_pdf_text(path):
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        
        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        content = await loop.run_in_executor(None, extract_pdf_text, file_path)
        return content
        
    except ImportError:
        # Fallback: Try to read as text (will likely fail for binary PDFs)
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = await f.read()
            return f"Note: PDF processing library not available. Raw content: {content[:1000]}..."
        except:
            return "Error: Could not process PDF file. Please install PyPDF2 library."

async def process_word_file(file_path: Path) -> str:
    """Process .docx/.doc files using python-docx"""
    try:
        from docx import Document
        
        def extract_docx_text(path):
            doc = Document(path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        
        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        content = await loop.run_in_executor(None, extract_docx_text, file_path)
        return content
        
    except ImportError:
        # Fallback: Try to read as text (will likely fail for binary docs)
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = await f.read()
            return f"Note: Word processing library not available. Raw content: {content[:1000]}..."
        except:
            return "Error: Could not process Word file. Please install python-docx library."

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    # Remove excessive whitespace
    lines = [line.strip() for line in text.split('\n')]
    lines = [line for line in lines if line]  # Remove empty lines
    
    # Join with single newlines
    cleaned = '\n'.join(lines)
    
    # Limit length if too long
    max_length = 50000  # 50KB text limit
    if len(cleaned) > max_length:
        cleaned = cleaned[:max_length] + "\n\n[Content truncated due to length...]"
    
    return cleaned